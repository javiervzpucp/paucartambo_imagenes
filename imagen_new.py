import os
import streamlit as st
from dotenv import load_dotenv
import pandas as pd
from openai import OpenAI
from datetime import datetime
from docx import Document
from docx.shared import Inches
import requests
from PIL import Image
from io import BytesIO
import tempfile

# Cargar las variables de entorno desde el archivo .env
load_dotenv()
openai_api_key = st.secrets["OPENAI_API_KEY"]
client = OpenAI(api_key=openai_api_key)

# Rutas de archivos CSV
dataset_path = "imagenes/imagenes.csv"
new_dataset_path = "imagenes/nuevas_descripciones.csv"

# Cargar o inicializar los DataFrames
df = pd.read_csv(dataset_path, delimiter=';', encoding='ISO-8859-1')
if os.path.exists(new_dataset_path):
    new_df = pd.read_csv(new_dataset_path, delimiter=';', encoding='ISO-8859-1')
else:
    new_df = pd.DataFrame(columns=["imagen", "descripcion", "generated_description", "keywords", "fecha"])

# Prompts para descripciones y palabras clave
describe_system_prompt = '''
Eres un sistema especializado en generar descripciones breves y precisas para escenas culturales y eventos andinos, especialmente de la festividad de la Mamacha Carmen en Paucartambo. 
Describe la escena de manera objetiva y clara, utilizando un estilo conciso. Asegúrate de destacar únicamente los elementos visibles y relevantes. 
No incluyas adornos adicionales ni interpretaciones extensas. Devuelve solo el texto de la descripción, sin formato adicional.
'''

keyword_system_prompt = '''
Eres un sistema experto en generar palabras clave concisas y relevantes para describir imágenes relacionadas con eventos culturales andinos. 
Tu respuesta debe incluir únicamente un arreglo de cadenas en formato JSON. 
Por ejemplo: ["máscara", "altar", "devoción", "sincretismo"]. 
No incluyas explicaciones ni texto adicional, solo las palabras clave en este formato exacto.
'''


def describe_image(img_path, title, example_descriptions):
    prompt = f"{describe_system_prompt}\n\nEjemplos de descripciones previas:\n{example_descriptions}\n\nGenera una descripción para la siguiente imagen:\nTítulo: {title}"
    response = client.chat.completions.create(
        model="gpt-4-turbo",
        messages=[
            {"role": "system", "content": describe_system_prompt},
            {"role": "user", "content": prompt}
        ],
        max_tokens=300,
        temperature=0.2
    )
    return response.choices[0].message.content.strip()


def generate_keywords(description):
    prompt = f"{keyword_system_prompt}\n\nDescripción: {description}"
    response = client.chat.completions.create(
        model="gpt-4-turbo",
        messages=[
            {"role": "system", "content": keyword_system_prompt},
            {"role": "user", "content": prompt}
        ],
        max_tokens=100,
        temperature=0.2
    )
    response_text = response.choices[0].message.content.strip()

    try:
        keywords = eval(response_text)
        if isinstance(keywords, list):
            return keywords
        else:
            raise ValueError("La respuesta no es una lista válida.")
    except (SyntaxError, ValueError):
        st.error(f"Error al analizar las palabras clave generadas. Respuesta original: {response_text}")
        return []


def export_to_word(description, keywords, date, title, img_path):
    doc = Document()
    doc.add_heading("Resumen Cultural", level=1)
    doc.add_paragraph(f"Fecha: {date}")
    doc.add_paragraph(f"Título: {title}")
    doc.add_paragraph(f"Descripción: {description}")
    doc.add_paragraph(f"Palabras clave: {', '.join(keywords)}")
    try:
        doc.add_picture(img_path, width=Inches(5.0))
    except Exception as e:
        st.error(f"No se pudo agregar la imagen: {e}")
    file_path = "resumen_cultural.docx"
    doc.save(file_path)
    return file_path


def get_combined_examples(df):
    combined_examples = "Ejemplos de descripciones previas:\n\n"
    for _, row in df.iterrows():
        if pd.notna(row.get('generated_description')) and pd.notna(row.get('descripcion')):
            combined_examples += f"Título: {row['descripcion']}\nDescripción: {row['generated_description']}\n\n"
    return combined_examples


# Interfaz de Streamlit
st.title("Generador de Descripciones de Imágenes de Danzas de Paucartambo")

# Sidebar para mostrar historial y opciones
with st.sidebar:
    st.write("Opciones")
    if st.checkbox("Mostrar historial"):
        st.dataframe(new_df[["imagen", "descripcion", "generated_description", "keywords"]])

# Entrada de imagen y descripción
option = st.radio("Seleccione el método para proporcionar una imagen:", ("URL de imagen", "Subir imagen"))

if option == "URL de imagen":
    img_url = st.text_input("Ingrese la URL de la imagen")
    title = st.text_input("Ingrese un título o descripción breve de la imagen")

    if img_url:
        try:
            response = requests.get(img_url, stream=True)
            response.raise_for_status()
            image = Image.open(BytesIO(response.content))
            st.image(image, caption="Imagen desde URL", use_column_width=True)
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_file:
                temp_file.write(response.content)
                img_path = temp_file.name
            example_descriptions = get_combined_examples(new_df)
            if st.button("Generar Descripción"):
                description = describe_image(img_path, title, example_descriptions)
                keywords = generate_keywords(description)
                st.write("Descripción generada:")
                st.write(description)
                st.write("Palabras clave generadas:")
                st.write(", ".join(keywords))
                file_path = export_to_word(description, keywords, datetime.now().strftime("%Y-%m-%d"), title, img_path)

                # Enlace para compartir en WhatsApp
                whatsapp_text = f"*Resumen*\n\n*Título:* {title}\n*Descripción:* {description}\n*Palabras clave:* {', '.join(keywords)}"
                whatsapp_link = f"https://wa.me/?text={requests.utils.quote(whatsapp_text)}"
                st.markdown(f"[Compartir en WhatsApp]({whatsapp_link})", unsafe_allow_html=True)

                with open(file_path, "rb") as file:
                    st.download_button(
                        label="Descargar Texto Resumen",
                        data=file,
                        file_name="resumen_imagen.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        except Exception as e:
            st.error(f"Error al procesar la URL de la imagen: {e}")

else:
    uploaded_file = st.file_uploader("Cargue una imagen", type=["jpg", "jpeg", "png"])
    title = st.text_input("Ingrese un título o descripción breve de la imagen")

    if uploaded_file:
        try:
            image = Image.open(uploaded_file)
            st.image(image, caption="Imagen cargada", use_column_width=True)
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_file:
                image.save(temp_file.name)
                img_path = temp_file.name
            example_descriptions = get_combined_examples(new_df)
            if st.button("Generar Descripción"):
                description = describe_image(img_path, title, example_descriptions)
                keywords = generate_keywords(description)
                st.write("Descripción generada:")
                st.write(description)
                st.write("Palabras clave generadas:")
                st.write(", ".join(keywords))
                file_path = export_to_word(description, keywords, datetime.now().strftime("%Y-%m-%d"), title, img_path)

                # Enlace para compartir en WhatsApp
                whatsapp_text = f"*Resumen Cultural*\n\n*Título:* {title}\n*Descripción:* {description}\n*Palabras clave:* {', '.join(keywords)}"
                whatsapp_link = f"https://wa.me/?text={requests.utils.quote(whatsapp_text)}"
                st.markdown(f"[Compartir en WhatsApp]({whatsapp_link})", unsafe_allow_html=True)

                with open(file_path, "rb") as file:
                    st.download_button(
                        label="Descargar Texto Resumen",
                        data=file,
                        file_name="resumen_imagen.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        except Exception as e:
            st.error(f"Error al procesar la imagen cargada: {e}")
